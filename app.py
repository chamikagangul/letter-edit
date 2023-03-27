from flask import Flask, render_template, request, send_file
from docx import Document
from openpyxl import load_workbook
import os
import copy

import zipfile

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_letter', methods=['POST'])
def generate_letter():
    # Get the uploaded files
    letter_file = request.files['letter']
    excel_file = request.files['excel']

    # Load the Word document and Excel sheet
    letter_doc = Document(letter_file)
    excel_workbook = load_workbook(excel_file)

    # Get the employee data from the Excel sheet
    employees_sheet = excel_workbook.active
    employees_data = []
    for row in employees_sheet.iter_rows(min_row=2, values_only=True):
        employee = {
            'empNo': row[1],
            'nic': row[2],
            'name': row[3],
            'doj': row[4].strftime("%Y/%m/%d"),
            'designation': row[5],
            'address': row[6],
            'basic': row[7],
            'bra': row[8]
        }
        employees_data.append(employee)

    # Fill in the placeholders in the Word document with the employee data
    temp_files = []
    for employee in employees_data:
        print(letter_doc)
        letter_doc_copy = fill_placeholders(copy.deepcopy(letter_doc), employee)

        # Generate a new Word document with the updated data
        generated_doc = generate_new_doc(letter_doc_copy)

        # Save the new document to a temporary file
        temp_path = os.path.join(os.getcwd(), 'temp/temp_{}.docx'.format(employee['empNo']))
        generated_doc.save(temp_path)
        temp_files.append(temp_path)

    # Create a zip file containing the updated documents
    zip_path = os.path.join(os.getcwd(), 'temp/documents.zip')
    with zipfile.ZipFile(zip_path, 'w') as zip_file:
        for temp_file in temp_files:
            zip_file.write(temp_file, os.path.basename(temp_file))

    # Send the zip file as a response to the user
    

    # Delete the temporary files
    for temp_file in temp_files:
        os.remove(temp_file)
    return send_file(zip_path, as_attachment=True)



def fill_placeholders(doc, employee):
    print(employee.items())
    # Replace the placeholders in the document with the employee data
    for p in doc.paragraphs:
        for key, value in employee.items():
            if '{' + key + '}' in p.text:
                p.text = p.text.replace('{' + key + '}', str(value))
    return doc

def generate_new_doc(doc):
    # Generate a new Word document with the updated data
    new_doc = Document()
    for element in doc.element.body:
        new_doc.element.body.append(element)
    return new_doc

if __name__ == '__main__':
    app.run(debug=True)
